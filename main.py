import streamlit as st
import os
import time
import requests
from newspaper import Article, Config
from docx import Document
from io import BytesIO
import google.generativeai as genai
from bs4 import BeautifulSoup

# API Key Setup
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
model = genai.GenerativeModel("models/gemini-1.5-flash-latest")

# App Title and UI
st.title("AI-Powered Content Brief Generator")
st.markdown("""
- 🔍 Scrapes top-performing content for your keyword  
- ✍️ Summarizes main themes with suggested headings and questions  
- 📋 Provides NLP-informed subtopics for optimized content  
""")

keyword = st.text_input("Enter a keyword:")
num_results = st.slider("Number of results", 1, 10, 3)

DATAFORSEO_LOGIN = os.getenv("DATAFORSEO_LOGIN")
DATAFORSEO_PASSWORD = os.getenv("DATAFORSEO_PASSWORD")

# Get SERP URLs
def get_serp_results(keyword, num_results):
    endpoint = "https://api.dataforseo.com/v3/serp/google/organic/live/regular"
    payload = {"location_code": 2840, "language_code": "en", "keyword": keyword, "limit": num_results}
    try:
        response = requests.post(endpoint, auth=(DATAFORSEO_LOGIN, DATAFORSEO_PASSWORD), json=[payload], timeout=10)
        response.raise_for_status()
        data = response.json()
        if "tasks" not in data or not data["tasks"]:
            st.error("No tasks returned from DataForSEO API")
            return []
        items = data["tasks"][0]["result"][0]["items"]
        return [item["url"] for item in items if "url" in item][:num_results]
    except Exception as e:
        st.error(f"Error fetching SERP results: {e}")
        return []

# Scrape Article Text
def get_article_text(url):
    try:
        config = Config()
        config.request_timeout = 10
        config.headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
        article = Article(url, config=config)
        article.download()
        article.parse()
        if article.text:
            return article.text
    except Exception as e:
        st.warning(f"Newspaper failed for {url}: {e}")
    try:
        response = requests.get(url, timeout=10, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'})
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "lxml")
        paragraphs = soup.find_all("p")
        text = " ".join([p.get_text() for p in paragraphs])
        if text:
            return text
        st.warning(f"No text extracted from {url} with BeautifulSoup")
        return ""
    except requests.exceptions.RequestException as e:
        st.warning(f"Failed to access {url}: {e}")
        return ""

# Summarize and Enhance with Headings, Questions, and Subtopics
def summarize_and_enhance(text, url):
    prompt = f"""
    Summarize the key points and topics from the following article:\n\n{text[:3000]}
    - Use ## for H2 headings to structure main themes.
    - For each theme, list 2-3 target questions prefixed with 'Q:' that users might ask.
    - For each theme, suggest 2-3 NLP-informed subtopics (e.g., key terms or entities) prefixed with 'Subtopics:' to deepen content.
    Ensure the output is strictly formatted as requested.
    """
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"⚠️ Gemini error: {e}"

# Generate .docx Summary File
def generate_docx(summaries):
    doc = Document()
    doc.add_heading("AI-Powered Content Brief", 0)
    for url, content in summaries:
        doc.add_heading(url, level=1)
        doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Main App Logic
summaries = []

if keyword:
    with st.spinner("Generating content brief..."):
        st.info(f"Analyzing the top {num_results} search results for '{keyword}'...")
        urls = get_serp_results(keyword, num_results)

        for url in urls:
            if any(domain in url for domain in ['wfu.edu', 'life-global.org', 'wikipedia.org']):
                st.warning(f"Skipping restricted domain: {url}")
                continue

            st.write(f"🔗 Analyzing: {url}")
            content = get_article_text(url)
            if content:
                enhanced_summary = summarize_and_enhance(content, url)
                summaries.append((url, enhanced_summary))
                with st.expander(f"✅ Summary for: {url}", expanded=False):
                    st.write(enhanced_summary)

        st.success("✅ Finished summarizing!")

        if summaries:
            docx_file = generate_docx(summaries)
            st.download_button(
                label="📥 Download all summaries as .docx",
                data=docx_file,
                file_name="content_brief.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("No summaries available to download.")

# Fallback Test
if keyword and not summaries:
    st.warning("No summaries were generated. Running Gemini test prompt...")
    test_text = "Project management involves planning, executing, and controlling a team to achieve specific goals."
    test_prompt = f"""
    Summarize the key points and topics from the following:\n\n{test_text}
    - Use ## for H2 headings to structure main themes.
    - For each theme, list 2-3 target questions prefixed with 'Q:' that users might ask.
    - For each theme, suggest 2-3 NLP-informed subtopics (e.g., key terms or entities) prefixed with 'Subtopics:' to deepen content.
    Ensure the output is strictly formatted as requested.
    """
    try:
        response = model.generate_content(test_prompt)
        st.markdown("### 🔍 Gemini test summary:")
        st.write(response.text)
    except Exception as e:
        st.error(f"⚠️ Gemini test failed: {repr(e)}")
